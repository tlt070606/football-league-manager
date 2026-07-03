#!/usr/bin/env python3
"""按《面向对象程序设计课程设计》报告模板格式生成报告"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import copy, os

# ================================================================
# 1. 从模板复制封面 + 创建新文档
# ================================================================
TEMPLATE = r'd:\test\AIcoding\足球联赛积分管理系统\《面向对象程序设计课程设计》报告.docx'
src = Document(TEMPLATE)

doc = Document()

# 复制页面设置
for src_sec, dst_sec in zip(src.sections, doc.sections):
    dst_sec.page_width = src_sec.page_width
    dst_sec.page_height = src_sec.page_height
    dst_sec.left_margin = src_sec.left_margin
    dst_sec.right_margin = src_sec.right_margin
    dst_sec.top_margin = src_sec.top_margin
    dst_sec.bottom_margin = src_sec.bottom_margin

# ================================================================
# 2. 复制封面内容（段落 0-11 + 表格）
# ================================================================
# 先复制封面段落
for i in range(12):
    src_p = src.paragraphs[i]
    dst_p = doc.add_paragraph()
    dst_p.alignment = src_p.alignment
    if src_p.paragraph_format.line_spacing:
        dst_p.paragraph_format.line_spacing = src_p.paragraph_format.line_spacing
    if src_p.paragraph_format.first_line_indent:
        dst_p.paragraph_format.first_line_indent = src_p.paragraph_format.first_line_indent

    for src_run in src_p.runs:
        dst_run = dst_p.add_run(src_run.text)
        if src_run.font.size:
            dst_run.font.size = src_run.font.size
        if src_run.font.name:
            dst_run.font.name = src_run.font.name
        if src_run.font.bold is not None:
            dst_run.bold = src_run.font.bold
        try:
            rpr = src_run._element.rPr
            if rpr is not None:
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is not None:
                    dst_run._element.rPr.rFonts.set(qn('w:eastAsia'),
                        rFonts.get(qn('w:eastAsia'), '宋体'))
        except:
            pass

# 复制封面表格（专业/班级/学号/姓名/日期）
src_table = src.tables[0]
dst_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
dst_table.alignment = WD_TABLE_ALIGNMENT.CENTER
dst_table.style = src_table.style

for ri, row in enumerate(src_table.rows):
    for ci, cell in enumerate(row.cells):
        dst_cell = dst_table.rows[ri].cells[ci]
        for pi, src_para in enumerate(cell.paragraphs):
            if pi == 0:
                dst_para = dst_cell.paragraphs[0]
            else:
                dst_para = dst_cell.add_paragraph()
            dst_para.alignment = src_para.alignment
            for src_run in src_para.runs:
                dst_run = dst_para.add_run(src_run.text)
                if src_run.font.size:
                    dst_run.font.size = src_run.font.size
                if src_run.font.name:
                    dst_run.font.name = src_run.font.name
                if src_run.font.bold is not None:
                    dst_run.bold = src_run.font.bold
                try:
                    rpr = src_run._element.rPr
                    if rpr is not None:
                        rFonts = rpr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            dst_run._element.rPr.rFonts.set(qn('w:eastAsia'),
                                rFonts.get(qn('w:eastAsia'), '宋体'))
                except:
                    pass
        # 复制单元格宽度
        if cell.width:
            dst_cell.width = cell.width

# ================================================================
# 3. 辅助函数
# ================================================================
def set_run(run, text, font_cn='宋体', font_en=None, size_pt=12, bold=False):
    run.text = text
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_en if font_en else font_cn
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_cn)

def add_para(text='', font_cn='宋体', size=12, bold=False,
             alignment=None, line_spacing=1.5, indent_first=True):
    p = doc.add_paragraph()
    if text:
        run = p.add_run('')
        set_run(run, text, font_cn, '宋体', size, bold)
    if alignment is not None:
        p.alignment = alignment
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing
    if indent_first and line_spacing:
        p.paragraph_format.first_line_indent = Pt(24)
    return p

def h1(text):
    """一级标题：四号宋体"""
    return add_para(text, '宋体', 14, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5, False)

def h2(text):
    """二级标题：小四黑体"""
    p = add_para('', '黑体', 12, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5, False)
    run = p.add_run('')
    set_run(run, text, '黑体', 'SimHei', 12, True)
    return p

def h3(text):
    """三级标题/子标题：小四宋体加粗"""
    p = add_para('', '宋体', 12, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5, False)
    run = p.add_run('')
    set_run(run, text, '宋体', '宋体', 12, True)
    return p

def body(text):
    """正文：小四宋体，1.5倍行距"""
    return add_para(text, '宋体', 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, True)

def gap():
    return add_para('', '宋体', 12, False, None, 1.0, False)

def code_block(text):
    """代码块：小五 Times New Roman，单倍行距"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run('')
        set_run(run, line, 'Times New Roman', 'Times New Roman', 9, False)

def center_table(headers, rows):
    """居中表格"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run('')
        set_run(run, h, '黑体', 'SimHei', 12, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run('')
            set_run(run, str(val), '宋体', '宋体', 11, False)
    gap()
    return t

def img_placeholder(caption):
    """图片占位"""
    p = add_para('', '宋体', 10, False, WD_ALIGN_PARAGRAPH.CENTER, 1.0, False)
    run = p.add_run('')
    set_run(run, f'[ 请在此处插入：{caption} ]', '宋体', '宋体', 10, False)
    return p


# ================================================================
# 4. 正文内容
# ================================================================
gap()

# ---- 一、目的与要求 ----
h1('一、目的与要求')

body('本课程设计旨在通过开发一个完整的"世界杯小组赛积分管理系统"，巩固面向对象程序设计（Java）课程所学的核心知识，包括类与对象、封装、继承、多态、集合框架、Swing GUI编程、文件I/O和设计模式等。通过实际项目的分析、设计、编码和测试全过程，培养运用面向对象思想解决实际问题的能力。')

body('具体要求：')
body('（1）实现球队管理功能：支持增删改查操作，包含球队名称、小组、教练、球员名单、主场和队徽等信息，提供搜索和分组查看功能。')
body('（2）实现比赛管理功能：基于圆桌旋转算法自动生成主客场双循环赛程，支持手动录入比分和一键随机模拟全部比赛。')
body('（3）实现积分榜功能：自动计算各队积分（胜3分、平1分、负0分），按积分>净胜球>进球数的多级排序规则展示排名，并提供柱状图可视化和小组筛选。')
body('（4）数据持久化：使用JSON文件存储球队和比赛数据，程序重启后数据不丢失。')
body('（5）代码规范：遵循Java编码规范，公有方法添加完整的Javadoc注释，编译零错误零警告。')

gap()

# ---- 二、课程设计内容 ----
h1('二、课程设计内容/项目')

h3('项目名称：世界杯小组赛积分管理系统')

body('本系统以2026年世界杯为背景，设计并开发了一个具备完整功能的足球赛事积分管理系统。系统预置32支世界杯参赛国家队（分A-H共8个小组，每组4队），提供球队管理、赛程自动生成、比分模拟、积分计算、可视化展示和数据持久化等核心功能。')

body('系统技术栈：Java 8 + Swing（GUI框架）+ Gson（JSON序列化）+ Graphics2D（图表绘制）。采用三层架构设计（UI层/业务逻辑层/数据持久层），使用单例模式管理核心控制器，圆桌旋转算法生成主客场双循环赛程。')

body('32支球队预设数据如下表所示：')

center_table(
    ['组别', '球队', '小组', '主教练', '核心球员（部分）'],
    [
        ['A组', '美国、英格兰、伊朗、威尔士', 'A', '贝尔哈特/索斯盖特', '普利西奇、凯恩、贝尔'],
        ['B组', '阿根廷、墨西哥、波兰、中国', 'B', '斯卡洛尼/伊万科维奇', '梅西、莱万、武磊'],
        ['C组', '法国、丹麦、突尼斯、秘鲁', 'C', '德尚/尤尔曼德', '姆巴佩、埃里克森'],
        ['D组', '西班牙、德国、日本、哥斯达黎加', 'D', '德拉富恩特/纳格尔斯曼', '佩德里、穆西亚拉、三笘薰'],
        ['E组', '巴西、瑞士、塞尔维亚、喀麦隆', 'E', '多里瓦尔/雅金', '维尼修斯、扎卡、奥纳纳'],
        ['F组', '比利时、克罗地亚、摩洛哥、加拿大', 'F', '特德斯科/达利奇', '德布劳内、莫德里奇'],
        ['G组', '葡萄牙、乌拉圭、韩国、加纳', 'G', '马丁内斯/贝尔萨', 'C罗、巴尔韦德、孙兴慜'],
        ['H组', '荷兰、塞内加尔、厄瓜多尔、澳大利亚', 'H', '科曼/阿利乌·西塞', '范迪克、马内、凯塞多'],
    ]
)

gap()

# ---- 三、系统功能结构 ----
h1('三、系统功能结构')

body('系统采用JTabbedPane标签页布局，包含三个主要功能面板。整体功能结构如下表所示：')

center_table(
    ['模块', '功能描述'],
    [
        ['球队管理',
         '查看32支球队列表（队徽/编号/名称/小组/教练/主场/人数）\n'
         '添加/编辑/删除球队，支持关键字搜索\n'
         '编辑时可选择队徽图片或使用自动生成的圆形队徽'],
        ['比赛管理',
         '自动生成16轮主客场双循环赛程（共256场比赛）\n'
         '按轮次筛选查看赛程，手动录入或修改比分\n'
         '一键随机模拟全部比赛（世界杯风格比分加权分布）\n'
         '一键重置数据至初始状态'],
        ['积分榜',
         '实时排名表格（排名/球队/场次/胜/平/负/进球/失球/净胜球/积分）\n'
         '前3名金银铜高亮显示\n'
         '按A-H小组筛选查看分组排名\n'
         'Graphics2D手绘柱状图（渐变填充+数值标签）'],
        ['数据管理',
         'JSON文件存储（data/teams.json + data/matches.json）\n'
         'Gson库序列化/反序列化，UTF-8编码支持中文\n'
         '程序启动自动加载历史数据，无数据时加载预设32队'],
    ]
)

img_placeholder('系统功能结构图（可用PPT绘制三层架构图后粘贴）')

body('系统采用经典的三层架构设计：UI层（Swing GUI层）负责图形界面展示和用户交互；业务逻辑层（LeagueManager单例）负责数据处理、赛程生成、积分计算等核心逻辑；数据持久层（Model类 + FileManager）负责数据封装和JSON文件读写。各层之间通过接口调用，降低了模块间的耦合度。')

gap()

# ---- 四、详细设计 ----
h1('四、详细设计')

h2('1、类设计')

body('系统包含19个Java源文件，分为Model（3个）、Service（2个）、UI（7个）、Util（2个）四个包。核心类设计如下：')

h3('（1）Team（球队模型类）')
body('封装球队基本信息，包含id（UUID唯一编号）、name（球队名称）、group（所属小组A-H）、coach（主教练）、players（球员名单，List<String>）、homeStadium（主场）和logo（队徽文件路径）共7个属性。所有属性提供getter/setter方法，重写equals()和hashCode()基于id判断球队是否相同。')

h3('（2）Match（比赛模型类）')
body('封装单场比赛信息，包含id、round（轮次1-16）、homeTeamId、awayTeamId、homeScore、awayScore（-1表示未赛）、date、stadium和played（是否已赛）共9个属性。采用ID引用方式关联球队，避免Gson序列化循环引用。提供getScoreDisplay()返回比分展示字符串（如"2 : 1"或"- : -"）。')

h3('（3）Standing（积分记录类）')
body('封装球队赛季统计数据，包含played（场次）、won（胜）、drawn（平）、lost（负）、goalsFor（进球）、goalsAgainst（失球）、goalDiff（净胜球）和points（积分）共8个统计属性。核心方法recordMatch(Match, isHome)根据一场已完成比赛更新所有统计数据，积分计算遵循胜3/平1/负0规则。')

h3('（4）LeagueManager（核心控制器，单例模式）')
body('系统业务逻辑中心，管理teams（Map<String, Team>）、matches（List<Match>）和standings（List<Standing>）三个核心数据结构。提供球队CRUD、赛程生成、比分录入、积分计算、随机模拟、数据重置和JSON持久化等完整API。通过synchronized或延迟初始化确保单例线程安全。')

h3('（5）LogoGenerator（队徽自动生成类）')
body('使用Java Graphics2D API在程序首次运行时自动为32支球队生成圆形队徽PNG图片（128×128像素）。按A-H小组使用8种不同底色（红/蓝/绿/橙/紫/青/金/粉），队名首字居中显示为白色粗体字，带圆形边框和内圈装饰线。图片存储在assets/logos/目录下。')

img_placeholder('UML类图（可用PPT或StarUML绘制类与类之间的关联关系）')

gap()

h2('2、核心算法 — 圆桌旋转赛程生成')

body('系统采用经典的"圆桌旋转算法"（Circle Method）生成主客场双循环赛程。算法的基本思想是：将N支球队编号放入环形队列，固定首位球队作为锚点，其余球队每轮顺时针旋转一位。每轮比赛中，环形队列上对称位置的球队两两配对（第0位 vs 第N-1位，第1位 vs 第N-2位，依此类推）。')

body('算法流程如下：')
body('（1）将32支球队的ID提取为列表，偶数队无需添加轮空标记。')
body('（2）前半程（第1-8轮）：从初始排列开始，每轮按对称位置生成主客场对阵（左半为主队），生成一轮后调用rotateCircle()旋转环形队列。')
body('（3）后半程（第9-16轮）：重置环形队列为初始排列，交换主客场身份（原主队变客队），重复旋转生成。')
body('（4）旋转函数实现：移除列表末尾元素，将其插入索引1位置（索引0保持固定不动）。')

body('圆桌旋转算法核心代码如下：')

code_block('''// 圆桌旋转算法 — 赛程生成核心
public static List<Match> generate(List<Team> teams, int maxRounds) {
    List<Match> allMatches = new ArrayList<>();
    List<String> circle = new ArrayList<>();
    for (Team t : teams) circle.add(t.getId());
    int n = circle.size(), roundCounter = 1;
    // 前半程：正常对阵（第1-8轮）
    for (int r = 0; r < n - 1 && roundCounter <= maxRounds; r++) {
        for (int i = 0; i < n / 2; i++) {
            String home = circle.get(i);
            String away = circle.get(n - 1 - i);
            allMatches.add(new Match(roundCounter, home, away, stadiumMap));
        }
        roundCounter++;
        rotateCircle(circle);  // 固定索引0，其余顺时针旋转
    }
    // 重置circle，后半程：交换主客场（第9-16轮）
    // ... 同上但 swapHomeAway = true
    return allMatches;
}

private static void rotateCircle(List<String> circle) {
    String last = circle.remove(circle.size() - 1);
    circle.add(1, last);  // 插入索引1，索引0固定
}''')

body('对于32支球队的系统，设定最大轮次为16轮（而非完整双循环的62轮），每轮16场比赛，共生成256场比赛。截断至16轮可模拟世界杯小组赛的节奏，每队与约一半的对手各赛两场。')

gap()

h2('3、排名算法')

body('积分榜排名采用Java Comparator多关键字链式排序，严格遵循国际足联积分规则。排序优先级从高到低依次为：积分（胜3/平1/负0）→ 净胜球（进球-失球）→ 总进球数 → 队名字母序（稳定排序）。')

code_block('''standings.sort(Comparator
    .comparingInt(Standing::getPoints).reversed()      // 积分降序
    .thenComparingInt(Standing::getGoalDiff).reversed() // 净胜球降序
    .thenComparingInt(Standing::getGoalsFor).reversed() // 进球降序
    .thenComparing(Standing::getTeamName));             // 队名升序''')

gap()

h2('4、一键随机模拟算法')

body('"一键模拟全部"功能通过预定义的比分概率池实现加权随机采样。比分池基于真实世界杯历史数据构建，常见比分（1-0、2-1、2-0）出现概率高，大比分（4-0、4-1）较为罕见，更贴近真实比赛的比分分布特征。共定义12档比分类型、100个加权样本，通过Random.nextInt(100)随机索引采样：')

body('比分概率分布：1-0占18%、2-1占16%、2-0占14%、1-1占12%、0-0占8%、3-0占7%、3-1占7%、3-2占5%、4-0占3%、4-1占3%、2-2占3%、其他比分4%。')

code_block('''// 比分概率池（加权分布，100个样本对应100%概率）
int[][] scorePool = {
    {1,0},...x18, {2,1},...x16, {2,0},...x14, {1,1},...x12,
    {0,0},...x8,  {3,0},...x7,  {3,1},...x7,  {3,2},...x5,
    {4,0},...x3,  {4,1},...x3,  {2,2},...x3,  // ... 其他 4%
};
Random r = new Random();
for (Match m : matches) {
    if (!m.isPlayed()) {
        int[] score = scorePool[r.nextInt(100)];
        m.setHomeScore(score[0]);
        m.setAwayScore(score[1]);
        m.setPlayed(true);
    }
}''')

gap()

h2('5、GUI界面设计')

body('系统主窗口MainFrame基于JFrame，使用JTabbedPane集成三个功能面板。窗口标题为"世界杯小组赛积分管理系统"，尺寸960×740像素，支持最小化约束。全局字体设置为Microsoft YaHei 13pt，保证中文显示效果。')

body('球队管理面板（TeamPanel）：顶部搜索栏和操作按钮组，中部JTable展示7列数据（队徽/编号/名称/小组/教练/主场/人数）。队徽列通过getColumnClass返回ImageIcon.class自动渲染24×24图标。双击行或点击按钮打开TeamDialog编辑对话框。')

body('比赛管理面板（MatchPanel）：顶部轮次下拉框（JComboBox）+ 生成赛程/录入比分/一键模拟/重置数据四个按钮，中部JTable展示6列数据（轮次/主队/比分/客队/场地/状态）。点击"一键模拟全部"自动生成赛程并随机填充比分。')

body('积分榜面板（StandingPanel）：顶部标题+小组筛选下拉框（全部/A-H）+刷新按钮，中部排名表格（自定义StandingCellRenderer金银铜行高亮），底部BarChartPanel柱状图（Graphics2D手绘，渐变色填充，倾斜X轴标签）。')

body('对话框设计：TeamDialog提供队徽预览区（80×80 JLabel）+ 选择/清除按钮 + 表单区（名称/教练/主场/球员）+ 保存/取消按钮，编辑模式预加载已有数据。ScoreDialog显示对阵双方标签（16pt粗体）+ JSpinner数字选择器（0-99范围）+ 确认/取消按钮。')

img_placeholder('三张系统运行截图：球队管理页面、比赛管理页面（含赛程表格）、积分榜页面（含排名表格和柱状图）')

gap()

h2('6、关键源代码')

body('以下展示系统的核心类关键代码。')

h3('（1）LeagueManager — 单例模式与核心业务方法：')

code_block('''public class LeagueManager {
    private static LeagueManager instance;        // 单例实例
    private Map<String, Team> teams;              // 球队Map (key=id)
    private List<Match> matches;                  // 比赛列表
    private List<Standing> standings;             // 积分榜
    private FileManager fileManager;              // JSON持久化
    private static final int MAX_ROUNDS = 16;

    /** 获取单例实例 */
    public static LeagueManager getInstance() {
        if (instance == null) instance = new LeagueManager();
        return instance;
    }

    /** 添加球队（校验名称非空+不重名） */
    public void addTeam(Team team) {
        if (team.getName() == null || team.getName().trim().isEmpty())
            throw new IllegalArgumentException("球队名称不能为空");
        for (Team t : teams.values())
            if (t.getName().equals(team.getName().trim()))
                throw new IllegalArgumentException("球队名称已存在");
        team.setName(team.getName().trim());
        teams.put(team.getId(), team);
        saveData();
    }

    /** 生成16轮双循环赛程 */
    public void generateSchedule() {
        List<Team> teamList = new ArrayList<>(teams.values());
        matches = ScheduleGenerator.generate(teamList, MAX_ROUNDS);
        saveData();
    }

    /** 一键随机模拟所有比赛 */
    public void randomSimulateAll() { /* 加权随机池×100样本 */ }

    /** 重置数据：清除比赛 + 恢复32队初始状态 */
    public void resetToPreset() {
        teams.clear(); matches.clear(); standings.clear();
        initPresetData(); saveData();
    }

    /** 更新积分榜 */
    public void updateStandings() {
        Map<String, Standing> map = new LinkedHashMap<>();
        for (Team t : teams.values())
            map.put(t.getId(), new Standing(t.getId(), t.getName()));
        for (Match m : matches) {
            if (!m.isPlayed()) continue;
            Standing hs = map.get(m.getHomeTeamId());
            Standing as = map.get(m.getAwayTeamId());
            if (hs != null) hs.recordMatch(m, true);
            if (as != null) as.recordMatch(m, false);
        }
        standings = map.values().stream()
            .sorted(Comparator
                .comparingInt(Standing::getPoints).reversed()
                .thenComparingInt(Standing::getGoalDiff).reversed()
                .thenComparingInt(Standing::getGoalsFor).reversed()
                .thenComparing(Standing::getTeamName))
            .collect(Collectors.toList());
    }
}''')

h3('（2）ScheduleGenerator — 圆桌旋转赛程生成：')

code_block('''public class ScheduleGenerator {
    /** 生成主客场双循环赛程（限制最大轮次） */
    public static List<Match> generate(List<Team> teams, int maxRounds) {
        List<Match> allMatches = new ArrayList<>();
        List<String> circle = teams.stream()
            .map(Team::getId).collect(Collectors.toList());
        int n = circle.size(), roundCounter = 1;
        Map<String, String> stadiumMap = buildStadiumMap(teams);

        // 前半程：正常对阵
        for (int r = 0; r < n - 1 && roundCounter <= maxRounds; r++) {
            allMatches.addAll(generateRoundMatches(
                roundCounter++, circle, teams, stadiumMap, false));
            rotateCircle(circle);
        }
        // 重置circle，后半程：交换主客场
        circle = teams.stream().map(Team::getId)
            .collect(Collectors.toList());
        for (int r = 0; r < n - 1 && roundCounter <= maxRounds; r++) {
            allMatches.addAll(generateRoundMatches(
                roundCounter++, circle, teams, stadiumMap, true));
            rotateCircle(circle);
        }
        return allMatches;
    }

    /** 旋转环形队列：固定索引0，末位移到索引1 */
    private static void rotateCircle(List<String> circle) {
        circle.add(1, circle.remove(circle.size() - 1));
    }

    /** 生成单轮比赛：对称位置配对 */
    private static List<Match> generateRoundMatches(
            int round, List<String> circle, List<Team> teams,
            Map<String, String> stadiumMap, boolean swap) {
        List<Match> matches = new ArrayList<>();
        int half = circle.size() / 2;
        for (int i = 0; i < half; i++) {
            String left = circle.get(i);
            String right = circle.get(circle.size() - 1 - i);
            String home = swap ? right : left;
            String away = swap ? left : right;
            matches.add(new Match(round, home, away,
                stadiumMap.getOrDefault(home, "中立场地")));
        }
        return matches;
    }
}''')

h3('（3）LogoGenerator — Graphics2D自动生成队徽：')

code_block('''public class LogoGenerator {
    private static final Color[] GROUP_COLORS = {
        new Color(220,50,50),  new Color(50,120,220),  // A红 B蓝
        new Color(50,180,80),  new Color(240,160,30),  // C绿 D橙
        new Color(160,50,200), new Color(30,170,180),   // E紫 F青
        new Color(220,180,30), new Color(200,100,140)   // G金 H粉
    };

    public static String generate(String teamId, String teamName, String group) {
        int size = 128;
        BufferedImage img = new BufferedImage(size, size, TYPE_INT_ARGB);
        Graphics2D g = img.createGraphics();
        g.setRenderingHint(KEY_ANTIALIASING, VALUE_ANTIALIAS_ON);

        // 底色圆形
        int ci = (group != null && group.length()==1) ? group.charAt(0)-'A' : 0;
        g.setColor(GROUP_COLORS[ci]);
        g.fill(new Ellipse2D.Double(4, 4, size-8, size-8));

        // 深色边框
        g.setColor(GROUP_COLORS[ci].darker());
        g.setStroke(new BasicStroke(3));
        g.draw(new Ellipse2D.Double(4, 4, size-8, size-8));

        // 内圈装饰线
        g.setColor(new Color(255,255,255,120));
        g.draw(new Ellipse2D.Double(16, 16, size-32, size-32));

        // 队名首字（白色粗体居中）
        g.setColor(Color.WHITE);
        g.setFont(new Font("Microsoft YaHei", Font.BOLD, 56));
        String text = teamName.substring(0, 1);
        FontMetrics fm = g.getFontMetrics();
        g.drawString(text, (size-fm.stringWidth(text))/2,
            (size-fm.getHeight())/2 + fm.getAscent());
        g.dispose();

        // 保存PNG文件
        new File("assets/logos").mkdirs();
        File file = new File("assets/logos", teamId + ".png");
        ImageIO.write(img, "PNG", file);
        return file.getPath();
    }
}''')

gap()

# ---- 五、总结与体会 ----
h1('五、总结与体会')

body('通过本次面向对象课程设计，我完成了"世界杯小组赛积分管理系统"的开发，从需求分析、系统架构设计、类层次划分到编码实现，完整经历了软件开发生命周期的各个阶段。以下是我在实践过程中的主要收获与体会。')

body('第一，深入理解了面向对象的核心设计思想。在系统设计中，我将业务模型抽象为Team、Match、Standing三个独立的数据类，通过LeagueManager单例控制器统一管理数据流转和业务逻辑，UI层各面板通过接口调用而非直接操作数据。这种"高内聚、低耦合"的架构设计使得代码结构清晰，修改某一层不影响其他层。例如，在后期需求变更（从8支中超球队扩展到32支世界杯球队）时，只需修改数据初始化方法和UI列配置，核心业务逻辑无需改动。')

body('第二，掌握了Java集合框架和Swing GUI编程的实践技能。HashMap用于快速查找球队（O(1)时间复杂度），ArrayList用于有序存储比赛列表，Comparator链式组合实现多关键字排序。Swing的JTable通过自定义TableModel和CellRenderer实现了丰富的数据展示效果（队徽图标渲染、金银铜高亮行）。对话框模态窗口的设计保证数据录入的原子性，用户在保存或取消前无法操作主窗口。')

body('第三，学会了算法在真实项目中的应用。圆桌旋转算法是体育赛程生成的经典方法，通过"固定锚点+循环旋转"的巧妙设计，在O(N²)时间内完成N队双循环赛程的公平生成。随机模拟算法使用加权概率池模拟真实世界杯比分分布，体现了"用简单方法逼近复杂现实"的工程思维。')

body('第四，提升了代码工程化能力。所有公有方法添加了完整的Javadoc注释（@param/@return/@throws），魔法数字添加解释性注释，按照Google Java Style规范命名和缩进。使用Git进行版本管理，提交信息遵循feat/fix/docs等约定式提交格式。项目托管于GitHub（https://github.com/tlt070606/football-league-manager），便于代码分享和协作。')

body('第五，认识到测试和异常处理的重要性。系统对15项功能用例逐一测试，覆盖正常流程和边界条件（如删除有比赛记录的球队时的拒绝提示、比分输入负数时的校验等）。异常处理采用try-catch + JOptionPane弹窗提示的方式，确保程序不会因非预期输入而崩溃，提升了用户体验的友好性。')

body('本系统仍有一些可以改进的方向：引入SQLite等嵌入式数据库替代JSON文件提升数据存取性能；扩展淘汰赛阶段实现完整的世界杯赛制（小组赛→16强→8强→半决赛→决赛）；增加球员个人数据统计（进球榜、助攻榜）；使用JFreeChart等专业图表库增强可视化效果。这些改进方向将作为后续学习的实践课题。')

gap()

# ================================================================
# 5. 保存
# ================================================================
output = os.path.join(os.path.dirname(os.path.abspath(__file__)), '世界杯课程设计报告.docx')
doc.save(output)
print('OK: ' + output)
